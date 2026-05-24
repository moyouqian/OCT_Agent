from pathlib import Path
from dataclasses import asdict, dataclass, field
from fnmatch import fnmatch
from typing import Any, Dict, Iterable, List, Optional
import re

from .chunking import build_child_chunks, build_parent_chunks
from .config import SelfRagConfig
from .document_parser import hash_file, parse_document
from .llm import build_embedding_backend
from .memory_store import ChromaMemoryStore
from .parent_store import SQLiteParentStore


SUPPORTED_EXTENSIONS = {".pdf", ".docx", ".txt", ".md"}
DEFAULT_EXCLUDE_GLOBS = (
    ".git/**",
    ".venv/**",
    ".pytest_cache/**",
    "__pycache__/**",
    "rag_store/**",
    "chroma_store/**",
    "bm25_store/**",
    "tmp_real_pdf_test/**",
    "tests/test_data/**",
)
BM25_FILENAMES = ("bm25_index.pkl", "bm25_docs.pkl", "bm25_meta.pkl")


@dataclass
class IngestCandidate:
    path: Path
    status: str = "pending"
    reason: str = ""
    duplicate_of: str = ""
    content_hash: str = ""


@dataclass
class ChunkQualityStats:
    documents: int = 0
    parent_chunks: int = 0
    child_chunks: int = 0
    assets: int = 0
    references: int = 0
    noise_blocks: int = 0
    furniture_blocks: int = 0
    reference_blocks: int = 0
    empty_parent_chunks: int = 0
    empty_child_chunks: int = 0
    oversize_parent_chunks: int = 0
    oversize_child_chunks: int = 0
    hard_limit_child_chunks: int = 0
    mojibake_warnings: int = 0
    warnings: List[str] = field(default_factory=list)

    def as_dict(self) -> Dict[str, Any]:
        return asdict(self)

    def add_children(self, title: str, children: List[Any], config: SelfRagConfig) -> None:
        self.child_chunks += len(children)
        for child in children:
            text = child.text.strip()
            if not text:
                self.empty_child_chunks += 1
            if estimated_child_tokens(child) > config.child_chunk_size:
                self.oversize_child_chunks += 1
            if (child.metadata or {}).get("split_reason") == "hard_limit":
                self.hard_limit_child_chunks += 1
            add_mojibake_warning(self, title, child.section_path, text, config.mojibake_extra_chars)


def extract_docx_text(file_path: str) -> str:
    from docx import Document

    doc = Document(file_path)
    parts = []
    for para in doc.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name.startswith("Heading"):
            parts.append(f"\n## {text}\n")
        else:
            parts.append(text)
    for table in doc.tables:
        for row in table.rows:
            cells = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if cells:
                parts.append("\t".join(cells))
    return "\n".join(parts).strip()


def ingest_file(
    file_path: str,
    config: Optional[SelfRagConfig] = None,
    embedder=None,
    store: Optional[ChromaMemoryStore] = None,
    parent_store: Optional[SQLiteParentStore] = None,
    rebuild_bm25: bool = True,
) -> Dict[str, Any]:
    config = config or SelfRagConfig()
    embedder = embedder or build_embedding_backend(config)
    store = store or ChromaMemoryStore(config.chroma_dir, config.collection_name)
    parent_store = parent_store or SQLiteParentStore(config.sqlite_path)

    path = Path(file_path)
    source_path = str(path.resolve())
    content_hash = hash_file(path)
    existing = parent_store.get_document_by_source_path(source_path)
    if existing and existing.get("content_hash") == content_hash:
        return {
            "skipped": 1,
            "total_parents": 0,
            "total_children": 0,
            "ingested_children": 0,
            "assets": 0,
            "references": 0,
            "noise_blocks": 0,
            "furniture_blocks": 0,
            "reference_blocks": 0,
            "grobid_status": "skipped",
        }

    if not existing:
        dup = parent_store.get_document_by_content_hash(content_hash)
        if dup and dup.get("source_path") != source_path:
            return {
                "skipped": 1,
                "total_parents": 0,
                "total_children": 0,
                "ingested_children": 0,
                "assets": 0,
                "references": 0,
                "noise_blocks": 0,
                "furniture_blocks": 0,
                "reference_blocks": 0,
                "grobid_status": "skipped",
                "duplicate_of": dup.get("source_path"),
            }

    if existing:
        child_ids = parent_store.list_child_ids_for_doc(existing["doc_id"])
        store.delete_ids(child_ids)
        parent_store.delete_document(existing["doc_id"])

    parsed = parse_document(str(path), config)
    parents, assets = build_parent_chunks(parsed, config)
    quality = analyze_chunks(parsed.title, parents, [], config)
    asset_counts = {}
    for asset in assets:
        if asset.parent_id:
            asset_counts[asset.parent_id] = asset_counts.get(asset.parent_id, 0) + 1
    for parent in parents:
        parent.metadata["asset_count"] = asset_counts.get(parent.parent_id, 0)

    ingested_children = 0
    total_children = 0
    stored_child_ids: List[str] = []
    try:
        parent_store.upsert_document(
            {
                "doc_id": parsed.doc_id,
                "source_path": parsed.source_path,
                "source_type": parsed.source_type,
                "file_type": parsed.file_type,
                "title": parsed.title,
                "authors": parsed.authors,
                "year": parsed.year,
                "content_hash": parsed.content_hash,
                "parser": parsed.parser,
                "parser_version": parsed.parser_version,
                "metadata": parsed.metadata,
            }
        )

        for parent in parents:
            parent_store.upsert_parent(asdict(parent))
            children = build_child_chunks(parent, config)
            quality.add_children(parsed.title, children, config)
            total_children += len(children)
            texts = [child.embedding_text for child in children]
            if hasattr(embedder, "get_embeddings"):
                embeddings = embedder.get_embeddings(texts)
            else:
                embeddings = [embedder.get_embedding(text) for text in texts]
            for child, embedding in zip(children, embeddings):
                store.store_child_chunk(
                    child.child_id,
                    child.embedding_text,
                    embedding,
                    child.metadata,
                )
                stored_child_ids.append(child.child_id)
                child_record = asdict(child)
                child_record["chroma_collection"] = config.collection_name
                parent_store.upsert_child(child_record)
                ingested_children += 1

        for asset in assets:
            parent_store.upsert_asset(asdict(asset))

        for reference in getattr(parsed, "references", []):
            parent_store.upsert_reference(asdict(reference))

        if rebuild_bm25:
            from .retrieval import rebuild_bm25_files

            rebuild_bm25_files(config, parent_store)
    except Exception:
        store.delete_ids(stored_child_ids)
        parent_store.delete_document(parsed.doc_id)
        raise

    return {
        "skipped": 0,
        "total_parents": len(parents),
        "total_children": total_children,
        "ingested_children": ingested_children,
        "assets": len(assets),
        "references": len(getattr(parsed, "references", [])),
        "noise_blocks": count_role(parsed.blocks, "noise"),
        "furniture_blocks": count_role(parsed.blocks, "furniture"),
        "reference_blocks": count_role(parsed.blocks, "reference"),
        "grobid_status": parsed.metadata.get("grobid", {}).get("status", ""),
        "empty_parent_chunks": quality.empty_parent_chunks,
        "empty_child_chunks": quality.empty_child_chunks,
        "oversize_parent_chunks": quality.oversize_parent_chunks,
        "oversize_child_chunks": quality.oversize_child_chunks,
        "hard_limit_child_chunks": quality.hard_limit_child_chunks,
        "mojibake_warnings": quality.mojibake_warnings,
    }


def ingest_many(paths: Iterable[str], config: Optional[SelfRagConfig] = None) -> Dict[str, Dict[str, Any]]:
    config = config or SelfRagConfig()
    embedder = build_embedding_backend(config)
    store = ChromaMemoryStore(config.chroma_dir, config.collection_name)
    parent_store = SQLiteParentStore(config.sqlite_path)
    results = {}
    for file_path in paths:
        results[file_path] = ingest_file(
            file_path,
            config=config,
            embedder=embedder,
            store=store,
            parent_store=parent_store,
            rebuild_bm25=False,
        )
    if any(not result.get("skipped") for result in results.values()):
        from .retrieval import rebuild_bm25_files

        rebuild_bm25_files(config, parent_store)
    return results


def discover_ingest_candidates(
    dataset_dirs: Optional[Iterable[str]] = None,
    files: Optional[Iterable[str]] = None,
    include_globs: Optional[Iterable[str]] = None,
    exclude_globs: Optional[Iterable[str]] = None,
    recursive: bool = True,
) -> List[IngestCandidate]:
    """Find supported files and apply scope rules before batch ingestion."""
    include_rules = tuple(include_globs or ())
    exclude_rules = tuple(DEFAULT_EXCLUDE_GLOBS) + tuple(exclude_globs or ())
    raw_paths: List[Path] = []

    for file_name in files or []:
        raw_paths.append(Path(file_name))

    for dirname in dataset_dirs or []:
        root = Path(dirname)
        if not root.exists():
            raw_paths.append(root)
            continue
        iterator = root.rglob("*") if recursive else root.glob("*")
        raw_paths.extend(path for path in iterator if path.is_file())

    candidates: List[IngestCandidate] = []
    seen_paths = set()
    seen_hashes: Dict[str, Path] = {}
    for path in sorted(raw_paths, key=lambda item: str(item).lower()):
        resolved = path.resolve()
        if resolved in seen_paths:
            continue
        seen_paths.add(resolved)
        candidate = IngestCandidate(path=resolved)
        candidates.append(candidate)

        if not path.exists():
            candidate.status = "excluded"
            candidate.reason = "missing"
            continue
        if path.suffix.lower() not in SUPPORTED_EXTENSIONS:
            candidate.status = "excluded"
            candidate.reason = "unsupported_extension"
            continue
        if include_rules and not any(match_path(path, rule) for rule in include_rules):
            candidate.status = "excluded"
            candidate.reason = "not_matched_by_include"
            continue
        if any(match_path(path, rule) for rule in exclude_rules):
            candidate.status = "excluded"
            candidate.reason = "matched_exclude_rule"
            continue

        candidate.content_hash = hash_file(path)
        if candidate.content_hash in seen_hashes:
            candidate.status = "duplicate"
            candidate.reason = "duplicate_content_hash"
            candidate.duplicate_of = str(seen_hashes[candidate.content_hash])
            continue
        seen_hashes[candidate.content_hash] = resolved
        candidate.status = "pending"

    return candidates


def run_batch_ingestion(
    candidates: Iterable[IngestCandidate],
    config: Optional[SelfRagConfig] = None,
    dry_run: bool = False,
    reset: bool = False,
    skip_quality_failures: bool = False,
    progress=None,
) -> Dict[str, Any]:
    config = config or SelfRagConfig()
    selected = [candidate for candidate in candidates if candidate.status == "pending"]
    excluded = [candidate for candidate in candidates if candidate.status != "pending"]
    summary: Dict[str, Any] = {
        "dry_run": dry_run,
        "reset": reset,
        "storage": storage_summary(config),
        "candidate_files": len(selected),
        "excluded_files": len(excluded),
        "processed_files": 0,
        "skipped_files": 0,
        "quality_skipped_files": 0,
        "failed_files": 0,
        "total_parents": 0,
        "total_children": 0,
        "ingested_children": 0,
        "assets": 0,
        "references": 0,
        "noise_blocks": 0,
        "furniture_blocks": 0,
        "reference_blocks": 0,
        "quality": ChunkQualityStats().as_dict(),
        "excluded": [asdict(candidate) for candidate in excluded],
        "quality_skipped": [],
        "failures": [],
        "results": {},
    }

    if reset and not dry_run:
        reset_ingestion_stores(config)

    if dry_run:
        for index, candidate in enumerate(selected, start=1):
            if progress:
                progress(index, len(selected), candidate.path, "inspect")
            try:
                quality = inspect_candidate_chunks(candidate.path, config)
            except Exception as exc:
                summary["failed_files"] += 1
                summary["failures"].append({"file": str(candidate.path), "error": str(exc)})
                continue
            merge_quality(summary, quality)
            summary["processed_files"] += 1
            summary["results"][str(candidate.path)] = {"dry_run": 1, **quality.as_dict()}
        return summary

    embedder = build_embedding_backend(config)
    store = ChromaMemoryStore(config.chroma_dir, config.collection_name)
    parent_store = SQLiteParentStore(config.sqlite_path)
    for index, candidate in enumerate(selected, start=1):
        if progress:
            progress(index, len(selected), candidate.path, "ingest")
        if skip_quality_failures:
            try:
                quality = inspect_candidate_chunks(candidate.path, config)
            except Exception as exc:
                summary["failed_files"] += 1
                summary["failures"].append({"file": str(candidate.path), "error": str(exc)})
                continue
            if quality_has_blockers(quality):
                summary["quality_skipped_files"] += 1
                merge_quality(summary, quality)
                summary["quality_skipped"].append(
                    {
                        "file": str(candidate.path),
                        "quality": quality.as_dict(),
                    }
                )
                continue
        try:
            result = ingest_file(
                str(candidate.path),
                config=config,
                embedder=embedder,
                store=store,
                parent_store=parent_store,
                rebuild_bm25=False,
            )
        except Exception as exc:
            summary["failed_files"] += 1
            summary["failures"].append({"file": str(candidate.path), "error": str(exc)})
            continue
        summary["results"][str(candidate.path)] = result
        if result.get("skipped"):
            summary["skipped_files"] += 1
        else:
            summary["processed_files"] += 1
        for key in (
            "total_parents",
            "total_children",
            "ingested_children",
            "assets",
            "references",
            "noise_blocks",
            "furniture_blocks",
            "reference_blocks",
        ):
            summary[key] += int(result.get(key, 0))
        merge_result_quality(summary, result)
    if summary["processed_files"] > 0 or summary["failed_files"] > 0:
        from .retrieval import rebuild_bm25_files

        rebuild_bm25_files(config, parent_store)
    return summary


def quality_has_blockers(quality: ChunkQualityStats) -> bool:
    return any(
        (
            quality.empty_parent_chunks,
            quality.empty_child_chunks,
            quality.oversize_parent_chunks,
            quality.oversize_child_chunks,
            quality.hard_limit_child_chunks,
            quality.mojibake_warnings,
        )
    )


def inspect_candidate_chunks(path: Path, config: SelfRagConfig) -> ChunkQualityStats:
    parsed = parse_document(str(path), config)
    parents, assets = build_parent_chunks(parsed, config)
    quality = analyze_chunks(parsed.title, parents, [], config)
    quality.assets = len(assets)
    quality.references = len(getattr(parsed, "references", []))
    quality.noise_blocks = count_role(parsed.blocks, "noise")
    quality.furniture_blocks = count_role(parsed.blocks, "furniture")
    quality.reference_blocks = count_role(parsed.blocks, "reference")
    for parent in parents:
        quality.add_children(parsed.title, build_child_chunks(parent, config), config)
    return quality


def analyze_chunks(title: str, parents: List[Any], children: List[Any], config: SelfRagConfig) -> ChunkQualityStats:
    quality = ChunkQualityStats(documents=1, parent_chunks=len(parents), child_chunks=len(children))
    for parent in parents:
        text = parent.text.strip()
        if not text:
            quality.empty_parent_chunks += 1
        if estimated_parent_tokens(parent) > config.parent_chunk_max_tokens:
            quality.oversize_parent_chunks += 1
        add_mojibake_warning(quality, title, parent.section_path, text, config.mojibake_extra_chars)
    quality.add_children(title, children, config)
    return quality


def estimated_parent_tokens(parent: Any) -> int:
    return int((parent.metadata or {}).get("estimated_tokens", 0))


def estimated_child_tokens(child: Any) -> int:
    return int((child.metadata or {}).get("estimated_tokens", 0))


def add_mojibake_warning(quality: ChunkQualityStats, title: str, section: str, text: str, extra_chars: str = "") -> None:
    if not text:
        return
    count = mojibake_score(text, extra_chars)
    if count >= 20 or count / max(len(text), 1) >= 0.02:
        quality.mojibake_warnings += 1
        if len(quality.warnings) < 20:
            quality.warnings.append(f"Possible mojibake in {title} / {section}")


KNOWN_MOJIBAKE_CHARS = "\u951f\u9286\u934f\u7b97\u9435\u701b\u9a9e\u6e2d\u5a09\u6434\u677b\u59af"
_CID_RE = re.compile(r"\(cid:\d+\)")


def mojibake_score(text: str, extra_chars: str = "") -> int:
    if not text:
        return 0
    score = 0
    score += text.count("\ufffd")  # U+FFFD replacement character
    score += len(_CID_RE.findall(text))  # leftover (cid:N) glyph fallbacks
    score += sum(1 for ch in text if "" <= ch <= "")  # private use area
    # C0/C1 control characters (except common whitespace)
    score += sum(1 for ch in text if (ord(ch) < 0x20 and ch not in "\n\t\r") or 0x80 <= ord(ch) <= 0x9f)
    # Known recurring mojibake glyphs from broken CJK encodings.
    # Note: legitimate Hanzi mangled into *other* valid Hanzi outside this table
    # cannot be detected heuristically \u2014 that is an inherent limitation.
    known = KNOWN_MOJIBAKE_CHARS + (extra_chars or "")
    score += sum(text.count(ch) for ch in known)
    return score


def count_role(blocks: Iterable[Any], role: str) -> int:
    return sum(1 for block in blocks if getattr(block, "role", "body") == role)


def merge_quality(summary: Dict[str, Any], quality: ChunkQualityStats) -> None:
    quality_dict = summary["quality"]
    for key, value in quality.as_dict().items():
        if key == "warnings":
            quality_dict[key].extend(value)
            quality_dict[key] = quality_dict[key][:50]
        else:
            quality_dict[key] += value
    summary["total_parents"] += quality.parent_chunks
    summary["total_children"] += quality.child_chunks
    summary["assets"] += quality.assets
    summary["references"] += quality.references
    summary["noise_blocks"] += quality.noise_blocks
    summary["furniture_blocks"] += quality.furniture_blocks
    summary["reference_blocks"] += quality.reference_blocks


def merge_result_quality(summary: Dict[str, Any], result: Dict[str, Any]) -> None:
    quality_dict = summary["quality"]
    if result.get("skipped"):
        return
    quality_dict["documents"] += 1
    quality_dict["parent_chunks"] += int(result.get("total_parents", 0))
    quality_dict["child_chunks"] += int(result.get("total_children", 0))
    quality_dict["assets"] += int(result.get("assets", 0))
    quality_dict["references"] += int(result.get("references", 0))
    quality_dict["noise_blocks"] += int(result.get("noise_blocks", 0))
    quality_dict["furniture_blocks"] += int(result.get("furniture_blocks", 0))
    quality_dict["reference_blocks"] += int(result.get("reference_blocks", 0))
    for key in (
        "empty_parent_chunks",
        "empty_child_chunks",
        "oversize_parent_chunks",
        "oversize_child_chunks",
        "hard_limit_child_chunks",
        "mojibake_warnings",
    ):
        quality_dict[key] += int(result.get(key, 0))


def match_path(path: Path, pattern: str) -> bool:
    normalized = path.as_posix()
    name = path.name
    if fnmatch(normalized, pattern) or fnmatch(name, pattern):
        return True
    if "/" in pattern and not Path(pattern).is_absolute():
        return fnmatch(normalized, f"*/{pattern}")
    return False


def storage_summary(config: SelfRagConfig) -> Dict[str, str]:
    return {
        "sqlite_path": config.sqlite_path,
        "chroma_dir": config.chroma_dir,
        "collection_name": config.collection_name,
        "bm25_dir": config.bm25_dir,
        "parsed_dir": config.parsed_dir,
        "assets_dir": config.assets_dir,
    }


def reset_ingestion_stores(config: SelfRagConfig) -> None:
    SQLiteParentStore(config.sqlite_path).reset()
    ChromaMemoryStore(config.chroma_dir, config.collection_name).reset()
    bm25_dir = Path(config.bm25_dir)
    for filename in BM25_FILENAMES:
        path = bm25_dir / filename
        if path.exists():
            path.unlink()


def format_batch_report(summary: Dict[str, Any]) -> str:
    quality = summary["quality"]
    lines = [
        "Ingestion report",
        "",
        f"- mode: {'dry-run' if summary['dry_run'] else 'ingest'}",
        f"- reset: {summary['reset']}",
        f"- candidate files: {summary['candidate_files']}",
        f"- excluded files: {summary['excluded_files']}",
        f"- processed files: {summary['processed_files']}",
        f"- skipped unchanged files: {summary['skipped_files']}",
        f"- skipped by quality gate: {summary.get('quality_skipped_files', 0)}",
        f"- failed files: {summary['failed_files']}",
        f"- parent chunks: {summary['total_parents']}",
        f"- child chunks: {summary['total_children']}",
        f"- ingested children: {summary['ingested_children']}",
        f"- assets: {summary['assets']}",
        f"- references: {summary['references']}",
        f"- noise blocks: {summary['noise_blocks']}",
        f"- furniture blocks: {summary['furniture_blocks']}",
        f"- reference blocks: {summary['reference_blocks']}",
        "",
        "Storage",
    ]
    for key, value in summary["storage"].items():
        lines.append(f"- {key}: {value}")
    lines.extend(
        [
            "",
            "Quality gates",
            f"- empty parent chunks: {quality['empty_parent_chunks']}",
            f"- empty child chunks: {quality['empty_child_chunks']}",
            f"- oversize parent chunks: {quality['oversize_parent_chunks']}",
            f"- oversize child chunks: {quality['oversize_child_chunks']}",
            f"- hard-limit child chunks: {quality['hard_limit_child_chunks']}",
            f"- possible mojibake warnings: {quality['mojibake_warnings']}",
            f"- references: {quality['references']}",
            f"- noise blocks: {quality['noise_blocks']}",
            f"- furniture blocks: {quality['furniture_blocks']}",
            f"- reference blocks: {quality['reference_blocks']}",
        ]
    )
    if quality["warnings"]:
        lines.extend(["", "Warnings"])
        lines.extend(f"- {warning}" for warning in quality["warnings"][:20])
    if summary["failures"]:
        lines.extend(["", "Failures"])
        lines.extend(f"- {item['file']}: {item['error']}" for item in summary["failures"][:20])
    if summary.get("quality_skipped"):
        lines.extend(["", "Quality-skipped files"])
        lines.extend(
            f"- {item['file']}: mojibake={item['quality']['mojibake_warnings']}, "
            f"empty_parent={item['quality']['empty_parent_chunks']}, "
            f"empty_child={item['quality']['empty_child_chunks']}, "
            f"oversize_parent={item['quality']['oversize_parent_chunks']}, "
            f"oversize_child={item['quality']['oversize_child_chunks']}, "
            f"hard_limit_child={item['quality']['hard_limit_child_chunks']}"
            for item in summary["quality_skipped"][:20]
        )
    duplicate_count = sum(1 for item in summary["excluded"] if item.get("status") == "duplicate")
    if duplicate_count:
        lines.extend(["", f"Duplicates skipped: {duplicate_count}"])
    return "\n".join(lines)
